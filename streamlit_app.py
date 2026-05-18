import streamlit as st
import requests
import io
import json
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

st.set_page_config(page_title="Generator Polityki Rachunkowosci", page_icon="📋", layout="wide")

ENTITY_FORM_LABELS = ["Sp. z o.o.", "Spolka akcyjna", "Spolka cywilna", "Spolka jawna",
                       "Spolka komandytowa", "Spolka kom.-akcyjna", "JDG", "Fundacja", "Stowarzyszenie"]
ENTITY_FORM_KEYS = ["sp_zoo", "sa", "sc", "sj", "sk", "ska", "jdg", "fundacja", "stowarzyszenie"]
ENTITY_FORM_FULL = {
    "sp_zoo": "Spolka z ograniczona odpowiedzialnoscia",
    "sa": "Spolka akcyjna", "sc": "Spolka cywilna", "sj": "Spolka jawna",
    "sk": "Spolka komandytowa", "ska": "Spolka komandytowo-akcyjna",
    "jdg": "Jednoosobowa dzialalnosc gospodarcza",
    "fundacja": "Fundacja", "stowarzyszenie": "Stowarzyszenie",
}
STEP_NAMES = ["Dane jednostki", "Ksiegi i Plan Kont", "Metody wyceny", "Koszty i RZiS",
              "Waluty obce", "Ochrona danych", "Polityki dodatkowe", "Podglad i eksport"]
ALL_CUR = ["EUR", "USD", "GBP", "CHF", "CZK", "SEK", "NOK", "DKK", "JPY", "CNY"]

if "step" not in st.session_state:
    st.session_state.step = 0
for k, v in dict(d_name="", d_form=0, d_nip="", d_krs="", d_regon="", d_addr="",
                  d_fys="01-01", d_fye="12-31", d_small=False, d_micro=False,
                  d_zpk="Wzorcowy plan kont", d_sn="", d_sv="", d_sp="",
                  d_dep="Metoda liniowa", d_thr=10000, d_iv="Cena nabycia", d_id="FIFO",
                  d_cm="Tylko Zespol 4 (uklad rodzajowy)", d_pl="Wariant porownawczy",
                  d_pc="Pelny koszt wytworzenia", d_oh="Klucz przychodowy",
                  d_fxs="Kurs sredni NBP", d_fxd="FIFO", d_hfx=False, d_cur=["EUR", "USD"],
                  d_dp="Elektroniczna i fizyczna", d_ay=5, d_bk="Codziennie", d_ac=True,
                  d_rp="", d_rev="Zasada memorialowa", d_ls="Wg przepisow bilansowych",
                  d_prov=True, d_dt=True, d_cf="Metoda posrednia",
                  d_adate=date.today(), d_edate=date.today(), d_ab="",
                  d_ksef=True, d_ksef_moment="Data wystawienia w KSeF",
                  d_ksef_korekty="Nota korygujaca w KSeF",
                  d_ksef_archiwum="Wylacznie w KSeF",
                  d_ksef_system="Zintegrowany z systemem FK").items():
    if k not in st.session_state:
        st.session_state[k] = v

def G(k):
    return st.session_state.get(k, "")

# ══════════════════════════════════════════════════════
# KRS API
# ══════════════════════════════════════════════════════

def fetch_krs_by_krs_nr(krs_nr):
    krs_clean = re.sub(r"[^0-9]", "", krs_nr).zfill(10)
    headers = {"Accept": "application/json", "User-Agent": "Mozilla/5.0 (compatible; PolitikaRachunkowosci/1.0)"}
    url = f"https://api-krs.ms.gov.pl/api/krs/OdpisAktualny/{krs_clean}"
    try:
        r = requests.get(url, params={"rejestr": "P", "format": "json"}, headers=headers, timeout=20)
        if r.status_code == 200:
            return _parse_odpis(r.json(), krs_clean)
        r2 = requests.get(url, params={"rejestr": "S", "format": "json"}, headers=headers, timeout=20)
        if r2.status_code == 200:
            return _parse_odpis(r2.json(), krs_clean)
    except requests.exceptions.ConnectionError:
        raise ConnectionError("Brak polaczenia z API KRS")
    except requests.exceptions.Timeout:
        raise TimeoutError("API KRS nie odpowiada")
    except Exception as e:
        raise RuntimeError(f"Blad API KRS: {e}")
    return None

def _parse_odpis(data, krs_nr=""):
    try:
        odpis = data.get("odpis", data)
        naglowek = odpis.get("naglowekA", {})
        dane = odpis.get("dane", {})
        dzial1 = dane.get("dzial1", {})
        dane_p = dzial1.get("danePodmiotu", {})
        nazwa = dane_p.get("nazwa", "")
        ident = dane_p.get("identyfikatory", {})
        nip_val = ident.get("nip", "")
        regon_raw = ident.get("regon", "")
        regon_val = regon_raw[:9] if regon_raw else ""
        forma = dane_p.get("formaPrawna", "")
        siedz_blok = dzial1.get("siedzibaIAdres", {})
        adres = siedz_blok.get("adres", {})
        ulica = adres.get("ulica", "")
        nr_domu = adres.get("nrDomu", "")
        nr_lok = adres.get("nrLokalu", "")
        kod = adres.get("kodPocztowy", "")
        miasto = adres.get("miejscowosc", "")
        siedziba = f"{ulica} {nr_domu}".strip()
        if nr_lok: siedziba += f"/{nr_lok}"
        if kod and miasto: siedziba += f", {kod} {miasto}"
        krs_val = naglowek.get("numerKRS", krs_nr)
        fl = forma.lower() if isinstance(forma, str) else ""
        forma_key = ("sp_zoo" if "ograniczon" in fl else "ska" if "komandytowo-akcyjn" in fl else
                     "sk" if "komandytow" in fl else "sa" if "akcyjn" in fl else
                     "sj" if "jawn" in fl else "fundacja" if "fundacj" in fl else
                     "stowarzyszenie" if "stowarzysz" in fl else "")
        dzial2 = dane.get("dzial2", {})
        sklad = dzial2.get("reprezentacja", {}).get("sklad", [])
        rep = ""
        if sklad:
            o = sklad[0]
            no = o.get("nazwisko", {})
            io2 = o.get("imiona", {})
            nz = no.get("nazwiskoICzlon", "") if isinstance(no, dict) else str(no)
            im = io2.get("imie", "") if isinstance(io2, dict) else str(io2)
            fn = o.get("funkcjaWOrganie", o.get("funkcja", ""))
            rep = f"{im} {nz}".strip()
            if fn: rep += f" - {fn}"
        return {"nazwa": nazwa, "siedziba": siedziba, "nip": nip_val, "krs": krs_val,
                "regon": regon_val, "forma_key": forma_key, "forma_prawna": forma, "rep": rep}
    except Exception:
        return None


# ══════════════════════════════════════════════════════
# ZPK GENERATOR — LOGIKA PLANU KONT 2026
# ══════════════════════════════════════════════════════

def generate_zpk(branza, typ_cit, wariant_rzis, skala, obsluga_aut, podmioty_powiazane):
    """Generuje Zakladowy Plan Kont na podstawie parametrow."""
    konta = []

    def add(kod, nazwa, typ, atr_pod, ksef=""):
        konta.append({"Kod_Konta": kod, "Nazwa_Konta": nazwa, "Typ": typ,
                       "Atrybut_Podatkowy": atr_pod, "Znacznik_KSeF": ksef})

    tp = ".TP" if podmioty_powiazane else ""

    # ── ZESPOL 0: Aktywa trwale ──
    add("010", "Srodki trwale", "Bilansowe", "-")
    add("011", "Wartosci niematerialne i prawne", "Bilansowe", "-")
    add("013", "Srodki trwale w budowie", "Bilansowe", "-")
    add("020", "Wartosci niematerialne i prawne - WNiP", "Bilansowe", "-")
    add("030", "Dlugoterminowe aktywa finansowe", "Bilansowe", "-")
    add("070", "Umorzenie srodkow trwalych", "Bilansowe", "-")
    add("071", "Umorzenie WNiP", "Bilansowe", "-")
    add("080", "Srodki trwale w budowie", "Bilansowe", "-")

    # ── ZESPOL 1: Srodki pieniezne ──
    add("100", "Kasa", "Bilansowe", "-")
    add("130", "Rachunki bankowe", "Bilansowe", "-")
    add("131", "Rachunek bankowy - biezacy PLN", "Bilansowe", "-")
    add("132", "Rachunek bankowy - walutowy", "Bilansowe", "-")
    add("135", "Rachunek VAT (split payment)", "Bilansowe", "-", "VAT_SPP")
    add("139", "Srodki pieniezne w drodze", "Bilansowe", "-")
    add("140", "Krotkoterminowe aktywa finansowe", "Bilansowe", "-")

    # ── ZESPOL 2: Rozrachunki ──
    add("200", "Rozrachunki z odbiorcami", "Bilansowe", "-", "FA_NAL")
    add("201", "Rozrachunki z dostawcami", "Bilansowe", "-", "FA_ZOB")
    add("220", "Rozrachunki publicznoprawne", "Bilansowe", "-")
    add("221", "Rozrachunki z US - VAT nalezny", "Bilansowe", "-", "VAT_NAL")
    add("222", "Rozrachunki z US - VAT naliczony", "Bilansowe", "-", "VAT_NAL")
    add("223", "Rozrachunki z US - CIT", "Bilansowe", "CIT")
    add("225", "Rozrachunki z US - PIT (pracownicy)", "Bilansowe", "-")
    add("229", "Rozrachunki z ZUS", "Bilansowe", "-")
    add("230", "Rozrachunki z pracownikami - wynagrodzenia", "Bilansowe", "-")
    add("234", "Rozrachunki z pracownikami - inne", "Bilansowe", "-")
    add("240", "Pozostale rozrachunki", "Bilansowe", "-")
    add("245", "Rozrachunki z wlascicielami/wspolnikami", "Bilansowe", "-")
    add("290", "Odpisy aktualizujace naleznosci", "Bilansowe", "NKUP")

    if podmioty_powiazane:
        add("200-TP", "Rozrachunki z odbiorcami - podmioty powiazane", "Bilansowe", "-", "FA_NAL_TP")
        add("201-TP", "Rozrachunki z dostawcami - podmioty powiazane", "Bilansowe", "-", "FA_ZOB_TP")

    # ── ZESPOL 3: Materialy i towary ──
    if branza in ["Produkcja", "Hybryda"]:
        add("310", "Materialy", "Bilansowe", "-")
        add("311", "Materialy na skladzie", "Bilansowe", "-")
        add("340", "Odchylenia od cen ewidencyjnych materialow", "Bilansowe", "-")

    if branza in ["Handel", "Hybryda"]:
        add("330", "Towary", "Bilansowe", "-")
        add("340", "Odchylenia od cen ewidencyjnych towarow", "Bilansowe", "-")

    add("300", "Rozliczenie zakupu", "Bilansowe", "-")

    # ── ZESPOL 4: Koszty rodzajowe ──
    if obsluga_aut:
        add("400", "Amortyzacja", "Wynikowe", "KUP")
        add("400-01", "Amortyzacja - KUP", "Wynikowe", "KUP")
        add("400-02", "Amortyzacja - NKUP (nadwyzka ponad limit)", "Wynikowe", "NKUP")
        add("401", "Zuzycie materialow i energii", "Wynikowe", "KUP")
        add("402", "Uslugi obce", "Wynikowe", "KUP")
        add("402-01", "Uslugi obce - KUP", "Wynikowe", "KUP")
        add("402-02", "Uslugi obce - NKUP (nadwyzka limit samochod)", "Wynikowe", "NKUP")
        add("403", "Podatki i oplaty", "Wynikowe", "KUP")
        add("404", "Wynagrodzenia", "Wynikowe", "KUP")
        add("405", "Ubezpieczenia spoleczne i inne swiadczenia", "Wynikowe", "KUP")
        add("409", "Pozostale koszty rodzajowe", "Wynikowe", "KUP")
    else:
        add("400", "Amortyzacja", "Wynikowe", "KUP")
        add("401", "Zuzycie materialow i energii", "Wynikowe", "KUP")
        add("402", "Uslugi obce", "Wynikowe", "KUP")
        add("403", "Podatki i oplaty", "Wynikowe", "KUP")
        add("404", "Wynagrodzenia", "Wynikowe", "KUP")
        add("405", "Ubezpieczenia spoleczne i inne swiadczenia", "Wynikowe", "KUP")
        add("409", "Pozostale koszty rodzajowe", "Wynikowe", "KUP")

    if podmioty_powiazane:
        add(f"402{tp}", "Uslugi obce - podmioty powiazane", "Wynikowe", "KUP")
        add(f"404{tp}", "Wynagrodzenia - podmioty powiazane", "Wynikowe", "KUP")

    # ── ZESPOL 5: Koszty wg typow dzialalnosci (kalkulacyjny) ──
    if wariant_rzis == "Kalkulacyjny" or branza in ["Produkcja", "Hybryda"]:
        add("501", "Koszty produkcji podstawowej", "Wynikowe", "KUP")
        add("520", "Koszty wydzialow", "Wynikowe", "KUP")
        add("527", "Koszty sprzedazy", "Wynikowe", "KUP")
        add("550", "Koszty ogolnego zarzadu", "Wynikowe", "KUP")
        add("580", "Rozliczenie kosztow dzialalnosci", "Wynikowe", "-")

        if branza in ["Produkcja", "Hybryda"]:
            add("530", "Koszty dzialalnosci pomocniczej", "Wynikowe", "KUP")

    # ── ZESPOL 6: Produkty i rozliczenia ──
    if branza in ["Produkcja", "Hybryda"]:
        add("601", "Wyroby gotowe", "Bilansowe", "-")
        add("602", "Polprodukty i produkcja w toku", "Bilansowe", "-")
        add("620", "Odchylenia od cen ewidencyjnych produktow", "Bilansowe", "-")

    add("640", "Rozliczenia miedzyokresowe kosztow czynne", "Bilansowe", "-")
    add("641", "Rozliczenia miedzyokresowe kosztow bierne", "Bilansowe", "-")

    # ── ZESPOL 7: Przychody ──
    add(f"700{tp}", "Przychody ze sprzedazy produktow", "Wynikowe", "Przychody_Op", "FA_PRZYCH")
    add(f"701{tp}", "Przychody ze sprzedazy uslug", "Wynikowe", "Przychody_Op", "FA_PRZYCH")

    if branza in ["Produkcja", "Hybryda"]:
        add("711", "Koszt wlasny sprzedazy produktow", "Wynikowe", "KUP")

    if branza in ["Handel", "Hybryda"]:
        add(f"730{tp}", "Przychody ze sprzedazy towarow", "Wynikowe", "Przychody_Op", "FA_PRZYCH")
        add("731", "Wartosc sprzedanych towarow w cenach zakupu", "Wynikowe", "KUP")

    add("740", "Przychody ze sprzedazy materialow", "Wynikowe", "Przychody_Op")
    add("741", "Wartosc sprzedanych materialow", "Wynikowe", "KUP")
    add("760", "Pozostale przychody operacyjne", "Wynikowe", "Przychody_Op")
    add("761", "Pozostale koszty operacyjne", "Wynikowe", "KUP")
    add("750", "Przychody finansowe", "Wynikowe", "Przychody_Kap")
    add("751", "Koszty finansowe", "Wynikowe", "KUP")
    add("770", "Zyski nadzwyczajne", "Wynikowe", "Przychody_Op")
    add("771", "Straty nadzwyczajne", "Wynikowe", "KUP")
    add("790", "Obroty wewnetrzne", "Wynikowe", "-")
    add("791", "Koszt obrotow wewnetrznych", "Wynikowe", "-")

    # ── ZESPOL 8: Kapital, rezerwy, wynik ──
    add("800", "Kapital zakladowy", "Bilansowe", "-")
    add("801", "Kapital zapasowy", "Bilansowe", "-")
    add("802", "Kapital rezerwowy", "Bilansowe", "-")
    add("803", "Kapital z aktualizacji wyceny", "Bilansowe", "-")
    add("810", "Zyski/straty z lat ubieglych", "Bilansowe", "-")
    add("820", "Rozliczenie wyniku finansowego", "Bilansowe", "-")

    if typ_cit == "Estonski":
        add("821", "Ukryte zyski (CIT estonski)", "Wynikowe", "NKUP")
        add("822", "Wydatki niezwiazane z dzialalnoscia (CIT estonski)", "Wynikowe", "NKUP")
        add("823", "Dochod z tyt. wydatkow niezwiazanych z dzialalnoscia", "Wynikowe", "NKUP")
        add("824", "Dochod z tyt. zmiany wartosci skladnikow majatku", "Wynikowe", "NKUP")

    add("840", "Rezerwy i rozliczenia miedzyokresowe przychodow", "Bilansowe", "-")
    add("841", "Rezerwa z tytulu odroczonego podatku dochodowego", "Bilansowe", "-")
    add("845", "Dotacje i subwencje", "Bilansowe", "-")
    add("850", "Fundusze specjalne (ZFSS)", "Bilansowe", "-")
    add("860", "Wynik finansowy", "Wynikowe", "-")
    add("870", "Obowiazkowe obciazenia wyniku finansowego - CIT", "Wynikowe", "CIT")

    return konta


def zpk_to_xlsx(konta):
    """Konwertuje liste kont na plik XLSX."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    except ImportError:
        return None

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Zakladowy Plan Kont"

    # Naglowki
    headers = ["Kod_Konta", "Nazwa_Konta", "Typ", "Atrybut_Podatkowy", "Znacznik_KSeF"]
    hfill = PatternFill(start_color="1B2A4A", end_color="1B2A4A", fill_type="solid")
    hfont = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
    thin = Side(style="thin", color="B0B0B0")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = hfont
        cell.fill = hfill
        cell.alignment = Alignment(horizontal="center")
        cell.border = border

    # Dane
    alt_fill = PatternFill(start_color="F2F6FA", end_color="F2F6FA", fill_type="solid")
    dfont = Font(name="Calibri", size=10)

    for i, konto in enumerate(konta, 2):
        vals = [konto["Kod_Konta"], konto["Nazwa_Konta"], konto["Typ"],
                konto["Atrybut_Podatkowy"], konto["Znacznik_KSeF"]]
        for col, v in enumerate(vals, 1):
            cell = ws.cell(row=i, column=col, value=v)
            cell.font = dfont
            cell.border = border
            if i % 2 == 0:
                cell.fill = alt_fill

    # Szerokosci kolumn
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 55
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 20
    ws.column_dimensions["E"].width = 18

    # Autofiltr
    ws.auto_filter.ref = f"A1:E{len(konta)+1}"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════
# SIDEBAR — DANE JEDNOSTKI + KRS
# ══════════════════════════════════════════════════════

with st.sidebar:
    st.header("Dane jednostki")
    krs_input = st.text_input("Numer KRS spolki", placeholder="0000640431")
    if st.button("Pobierz dane z KRS", use_container_width=True):
        if krs_input:
            with st.spinner("Pobieranie z API KRS..."):
                try:
                    krs_data = fetch_krs_by_krs_nr(krs_input)
                    if krs_data:
                        st.session_state["krs_data"] = krs_data
                        st.success("Dane pobrane z KRS!")
                    else:
                        st.error("Nie znaleziono.")
                except Exception as e:
                    st.error(f"Blad: {e}")
        else:
            st.warning("Wpisz numer KRS.")

    krs = st.session_state.get("krs_data", {})
    st.session_state.d_name = st.text_input("Nazwa spolki", value=krs.get("nazwa", G("d_name")))
    st.session_state.d_addr = st.text_input("Siedziba", value=krs.get("siedziba", G("d_addr")))
    st.session_state.d_nip = st.text_input("NIP", value=krs.get("nip", G("d_nip")))
    st.session_state.d_krs = st.text_input("Nr KRS", value=krs.get("krs", G("d_krs")))
    st.session_state.d_regon = st.text_input("REGON", value=krs.get("regon", G("d_regon")))
    if krs.get("forma_key") and krs["forma_key"] in ENTITY_FORM_KEYS:
        dfi = ENTITY_FORM_KEYS.index(krs["forma_key"])
    else:
        dfi = G("d_form") if isinstance(G("d_form"), int) else 0
    fv = st.selectbox("Forma prawna", ENTITY_FORM_LABELS, index=dfi)
    st.session_state.d_form = ENTITY_FORM_LABELS.index(fv) if fv in ENTITY_FORM_LABELS else 0
    if krs.get("rep") and not G("d_ab"):
        st.session_state.d_ab = krs["rep"]
    st.divider()
    st.subheader("Rok obrotowy")
    st.session_state.d_fys = st.text_input("Poczatek (MM-DD)", value=G("d_fys"))
    st.session_state.d_fye = st.text_input("Koniec (MM-DD)", value=G("d_fye"))
    st.session_state.d_small = st.checkbox("Jednostka mala (art. 3 ust. 1c)", value=G("d_small"))
    st.session_state.d_micro = st.checkbox("Jednostka mikro (art. 3 ust. 1a)", value=G("d_micro"))


# ══════════════════════════════════════════════════════
# DOCX GENERATION
# ══════════════════════════════════════════════════════

def gen_docx():
    """Generuje profesjonalnie sformatowany dokument Polityki Rachunkowości."""
    from docx.enum.table import WD_TABLE_ALIGNMENT

    # Polish quote characters via unicode escapes (safe in Python strings)
    LQ = "\u201E"  # „
    RQ = "\u201D"  # "

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)

    # Style normalny
    ns = doc.styles["Normal"]
    ns.font.name = "Calibri"
    ns.font.size = Pt(11)
    ns.paragraph_format.space_after = Pt(6)
    ns.paragraph_format.line_spacing = 1.35
    ns.paragraph_format.first_line_indent = Cm(0.5)

    # Style nagłówków
    for lv, (sz, cl) in {0: (18, "1A3C5E"), 1: (14, "2B5E8C"), 2: (12, "3B6B4F")}.items():
        h = doc.styles[f"Heading {lv+1}"]
        h.font.name = "Calibri"
        h.font.size = Pt(sz)
        h.font.bold = True
        h.font.color.rgb = RGBColor.from_string(cl)
        h.paragraph_format.space_before = Pt(22 if lv == 0 else 14)
        h.paragraph_format.space_after = Pt(10)
        h.paragraph_format.first_line_indent = Cm(0)
        h.paragraph_format.keep_with_next = True

    # Nagłówek strony
    hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Polityka Rachunkowo\u015bci \u2014 " + G("d_name"))
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(120, 120, 120)
    hr.font.italic = True
    hr.font.name = "Calibri"

    # Stopka z numeracją
    fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = fp.add_run("Strona ")
    rf.font.size = Pt(9)
    rf.font.color.rgb = RGBColor(120, 120, 120)
    rf.font.name = "Calibri"
    rp = fp.add_run()
    rp.font.size = Pt(9)
    rp.font.name = "Calibri"
    f1 = OxmlElement("w:fldChar")
    f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " PAGE "
    f2 = OxmlElement("w:fldChar")
    f2.set(qn("w:fldCharType"), "end")
    rp._r.append(f1)
    rp._r.append(it)
    rp._r.append(f2)

    # Pomocnicze funkcje
    def P(t, b=False, indent=True):
        pp = doc.add_paragraph()
        if not indent:
            pp.paragraph_format.first_line_indent = Cm(0)
        r = pp.add_run(t)
        r.bold = b
        r.font.name = "Calibri"
        return pp

    def PJ(t, b=False, indent=True):
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if not indent:
            pp.paragraph_format.first_line_indent = Cm(0)
        r = pp.add_run(t)
        r.bold = b
        r.font.name = "Calibri"
        return pp

    def PC(t, sz=11, b=False, i=False, c=None):
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.paragraph_format.first_line_indent = Cm(0)
        r = pp.add_run(t)
        r.font.size = Pt(sz)
        r.bold = b
        r.font.italic = i
        r.font.name = "Calibri"
        if c:
            r.font.color.rgb = RGBColor.from_string(c)

    def add_table(headers, rows, col_widths=None):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = False

        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.name = "Calibri"
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "1A3C5E")
            tcPr.append(shd)

        for ri, row in enumerate(rows):
            for i, val in enumerate(row):
                cell = t.rows[ri + 1].cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(str(val))
                r.font.size = Pt(10)
                r.font.name = "Calibri"
                if ri % 2 == 0:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "F2F6FA")
                    tcPr.append(shd)

        tbl = t._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        borders = OxmlElement("w:tblBorders")
        for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            b = OxmlElement("w:" + edge)
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), "4")
            b.set(qn("w:color"), "B0B0B0")
            borders.append(b)
        tblPr.append(borders)

        if col_widths:
            for ri in range(len(rows) + 1):
                for i, w in enumerate(col_widths):
                    t.rows[ri].cells[i].width = Cm(w)

        doc.add_paragraph()
        return t

    # Dane
    efi = G("d_form")
    efk = ENTITY_FORM_KEYS[efi] if isinstance(efi, int) and efi < len(ENTITY_FORM_KEYS) else ""
    efl = ENTITY_FORM_FULL.get(efk, "")
    efl_pl = {
        "Spolka z ograniczona odpowiedzialnoscia": "Spó\u0142ka z ograniczon\u0105 odpowiedzialno\u015bci\u0105",
        "Spolka akcyjna": "Spó\u0142ka akcyjna",
        "Spolka cywilna": "Spó\u0142ka cywilna",
        "Spolka jawna": "Spó\u0142ka jawna",
        "Spolka komandytowa": "Spó\u0142ka komandytowa",
        "Spolka komandytowo-akcyjna": "Spó\u0142ka komandytowo-akcyjna",
        "Jednoosobowa dzialalnosc gospodarcza": "Jednoosobowa dzia\u0142alno\u015b\u0107 gospodarcza",
        "Fundacja": "Fundacja",
        "Stowarzyszenie": "Stowarzyszenie",
    }.get(efl, efl)

    ad = G("d_adate"); ed = G("d_edate")
    ads = ad.strftime("%d.%m.%Y") if isinstance(ad, date) else str(ad)
    eds = ed.strftime("%d.%m.%Y") if isinstance(ed, date) else str(ed)
    thr = "{:,}".format(G("d_thr")).replace(",", " ")

    # ═══════ STRONA TYTUŁOWA ═══════
    for _ in range(4):
        doc.add_paragraph()
    PC("POLITYKA RACHUNKOWO\u015aCI", 26, True, c="1A3C5E")
    doc.add_paragraph()
    PC(G("d_name") or "[nazwa jednostki]", 16, True, c="2B5E8C")
    PC(efl_pl, 12, i=True, c="666666")
    doc.add_paragraph()
    PC("NIP: " + (G("d_nip") or "\u2014") + "   |   KRS: " + (G("d_krs") or "\u2014") + "   |   REGON: " + (G("d_regon") or "\u2014"), 10, c="666666")
    PC(G("d_addr") or "[adres siedziby]", 10, c="666666")
    for _ in range(2):
        doc.add_paragraph()
    PC("Opracowana na podstawie:", 10, c="666666")
    PC("Ustawy z dnia 29 wrze\u015bnia 1994 r. o rachunkowo\u015bci", 11, i=True, c="333333")
    PC("(t.j. Dz.U. z 2023 r. poz. 120 z pó\u017an. zm.)", 10, i=True, c="666666")
    PC("oraz Krajowych Standardów Rachunkowo\u015bci", 11, i=True, c="333333")
    for _ in range(3):
        doc.add_paragraph()
    PC("Obowi\u0105zuje od dnia: " + eds, 12, b=True, c="1A3C5E")
    PC("Data zatwierdzenia: " + ads, 11)
    doc.add_paragraph()
    PC("Zatwierdzi\u0142(a): " + (G("d_ab") or "__________________________"), 11)
    doc.add_page_break()

    # ═══════ SPIS TREŚCI ═══════
    doc.add_heading("Spis tre\u015bci", level=1)
    toc_items = [
        ("I.", "Postanowienia ogólne"),
        ("II.", "Zak\u0142adowy Plan Kont i prowadzenie ksi\u0105g rachunkowych"),
        ("III.", "Dowody ksi\u0119gowe i ich obieg"),
        ("IV.", "Metody wyceny aktywów i pasywów"),
        ("V.", "Ewidencja kosztów i wariant Rachunku Zysków i Strat"),
        ("VI.", "Operacje gospodarcze w walutach obcych"),
        ("VII.", "Inwentaryzacja"),
        ("VIII.", "System ochrony danych i ich zbiorów"),
        ("IX.", "Raportowanie elektroniczne (JPK, KSeF)"),
        ("X.", "Zasady dodatkowe i polityki szczególne"),
        ("XI.", "Postanowienia ko\u0144cowe"),
        ("", "Za\u0142\u0105czniki"),
    ]
    for num, title in toc_items:
        pp = doc.add_paragraph()
        pp.paragraph_format.first_line_indent = Cm(0)
        pp.paragraph_format.left_indent = Cm(0.5)
        pp.paragraph_format.space_after = Pt(4)
        r1 = pp.add_run((num + "  ") if num else "      ")
        r1.bold = True
        r1.font.size = Pt(11)
        r1.font.color.rgb = RGBColor.from_string("1A3C5E")
        r2 = pp.add_run(title)
        r2.font.size = Pt(11)
    doc.add_page_break()

    # ═══════ I. POSTANOWIENIA OGÓLNE ═══════
    doc.add_heading("I. Postanowienia ogólne", level=1)

    doc.add_heading("\u00a7 1. Podstawa prawna", level=2)
    PJ("1. Niniejsza Polityka Rachunkowo\u015bci (zwana dalej " + LQ + "Polityk\u0105" + RQ + ") zosta\u0142a opracowana na podstawie przepisów Ustawy z dnia 29 wrze\u015bnia 1994 r. o rachunkowo\u015bci (t.j. Dz.U. z 2023 r. poz. 120 z pó\u017an. zm.), zwanej dalej " + LQ + "Ustaw\u0105" + RQ + ", oraz Krajowych Standardów Rachunkowo\u015bci (KSR) wydawanych przez Komitet Standardów Rachunkowo\u015bci.")
    PJ("2. W sprawach nieuregulowanych Ustaw\u0105 oraz KSR jednostka stosuje odpowiednio Mi\u0119dzynarodowe Standardy Rachunkowo\u015bci (MSR) oraz Mi\u0119dzynarodowe Standardy Sprawozdawczo\u015bci Finansowej (MSSF).")
    PJ("3. Polityka okre\u015bla zasady i metody prowadzenia ksi\u0105g rachunkowych, wyceny aktywów i pasywów, ustalania wyniku finansowego oraz sporz\u0105dzania sprawozdania finansowego jednostki.")

    doc.add_heading("\u00a7 2. Dane identyfikacyjne jednostki", level=2)
    PJ("1. Polityka dotyczy jednostki: " + (G("d_name") or "[nazwa jednostki]") + ".")
    add_table(
        ["Pozycja", "Dane jednostki"],
        [
            ("Nazwa", G("d_name") or "\u2014"),
            ("Forma prawna", efl_pl or "\u2014"),
            ("NIP", G("d_nip") or "\u2014"),
            ("REGON", G("d_regon") or "\u2014"),
            ("KRS", G("d_krs") or "\u2014"),
            ("Siedziba", G("d_addr") or "\u2014"),
        ],
        col_widths=[5, 10]
    )

    doc.add_heading("\u00a7 3. Rok obrotowy", level=2)
    fys = "1 stycznia" if G("d_fys") == "01-01" else G("d_fys")
    fye = "31 grudnia" if G("d_fye") == "12-31" else G("d_fye")
    PJ("1. Rokiem obrotowym jednostki jest okres od " + fys + " do " + fye + " ka\u017cdego roku kalendarzowego.")
    PJ("2. Rok obrotowy dzieli si\u0119 na okresy sprawozdawcze obejmuj\u0105ce poszczególne miesi\u0105ce kalendarzowe.")
    PJ("3. Pierwszy rok obrotowy mo\u017ce by\u0107 d\u0142u\u017cszy ni\u017c 12 kolejnych miesi\u0119cy, jednak nie d\u0142u\u017cszy ni\u017c 18 miesi\u0119cy (art. 3 ust. 1 pkt 9 Ustawy).")

    doc.add_heading("\u00a7 4. J\u0119zyk i waluta", level=2)
    PJ("1. Ksi\u0119gi rachunkowe prowadzone s\u0105 w j\u0119zyku polskim.")
    PJ("2. Walut\u0105 prowadzenia ksi\u0105g jest z\u0142oty polski (PLN).")
    PJ("3. Warto\u015bci wyra\u017cone w walutach obcych przelicza si\u0119 na walut\u0119 polsk\u0105 wed\u0142ug zasad okre\u015blonych w Rozdziale VI niniejszej Polityki.")

    doc.add_heading("\u00a7 5. Status jednostki", level=2)
    if G("d_small"):
        PJ("1. Jednostka spe\u0142nia kryteria okre\u015blone w art. 3 ust. 1c Ustawy i klasyfikowana jest jako jednostka ma\u0142a. Korzysta z uproszcze\u0144 przewidzianych dla jednostek ma\u0142ych, w szczególno\u015bci w zakresie:")
        P("\u2022 sporz\u0105dzania uproszczonego bilansu i rachunku zysków i strat (Za\u0142\u0105cznik nr 5 do Ustawy),", indent=False)
        P("\u2022 zwolnienia z obowi\u0105zku sporz\u0105dzania zestawienia zmian w kapitale (funduszu) w\u0142asnym,", indent=False)
        P("\u2022 zwolnienia z obowi\u0105zku sporz\u0105dzania rachunku przep\u0142ywów pieni\u0119\u017cnych (o ile nie podlega badaniu).", indent=False)
    elif G("d_micro"):
        PJ("1. Jednostka spe\u0142nia kryteria okre\u015blone w art. 3 ust. 1a Ustawy i klasyfikowana jest jako jednostka mikro. Korzysta z maksymalnych uproszcze\u0144 przewidzianych dla jednostek mikro (Za\u0142\u0105cznik nr 4 do Ustawy).")
    else:
        PJ("1. Jednostka stosuje pe\u0142ne zasady rachunkowo\u015bci zgodnie z Ustaw\u0105, bez korzystania z uproszcze\u0144 przewidzianych dla jednostek ma\u0142ych lub mikro.")

    doc.add_heading("\u00a7 6. Odpowiedzialno\u015b\u0107", level=2)
    PJ("1. Za przestrzeganie zasad (polityki) rachunkowo\u015bci oraz prowadzenie ksi\u0105g rachunkowych odpowiedzialno\u015b\u0107 ponosi kierownik jednostki, zgodnie z art. 4 ust. 5 Ustawy.")
    PJ("2. Kierownik jednostki mo\u017ce powierzy\u0107 prowadzenie ksi\u0105g rachunkowych podmiotowi zewn\u0119trznemu, co nie zwalnia go z odpowiedzialno\u015bci okre\u015blonej w ust. 1.")

    # ═══════ II. ZPK I KSIĘGI ═══════
    doc.add_heading("II. Zak\u0142adowy Plan Kont i prowadzenie ksi\u0105g rachunkowych", level=1)

    doc.add_heading("\u00a7 7. Zak\u0142adowy Plan Kont", level=2)
    zpk = "wzorcowy plan kont" if "Wzorcowy" in G("d_zpk") else "indywidualnie opracowany plan kont, uwzgl\u0119dniaj\u0105cy specyfik\u0119 dzia\u0142alno\u015bci jednostki"
    PJ("1. Jednostka stosuje Zak\u0142adowy Plan Kont (ZPK) oparty o " + zpk + ", stanowi\u0105cy Za\u0142\u0105cznik nr 1 do niniejszej Polityki.")
    PJ("2. ZPK obejmuje wykaz kont ksi\u0119gi g\u0142ównej (syntetycznych) oraz kont ksi\u0105g pomocniczych (analitycznych) wraz z opisem ich przeznaczenia, zasad funkcjonowania oraz powi\u0105za\u0144 korespondencyjnych.")
    PJ("3. Zmiany w ZPK wprowadzane s\u0105 w trakcie roku obrotowego wy\u0142\u0105cznie w uzasadnionych przypadkach, z zachowaniem zasady ci\u0105g\u0142o\u015bci (art. 5 ust. 1 Ustawy). Zmiany dokumentuje si\u0119 aneksem do niniejszej Polityki.")

    doc.add_heading("\u00a7 8. Rodzaje ksi\u0105g rachunkowych", level=2)
    PJ("1. Ksi\u0119gi rachunkowe jednostki obejmuj\u0105, zgodnie z art. 13 Ustawy:")
    add_table(
        ["Lp.", "Rodzaj ksi\u0119gi", "Opis"],
        [
            ("1.", "Dziennik", "Zapis chronologiczny wszystkich zdarze\u0144, jakie nast\u0105pi\u0142y w okresie sprawozdawczym"),
            ("2.", "Ksi\u0119ga g\u0142ówna", "Konta syntetyczne \u2014 zapisy systematyczne zdarze\u0144 gospodarczych"),
            ("3.", "Ksi\u0119gi pomocnicze", "Konta analityczne uszczegó\u0142awiaj\u0105ce zapisy ksi\u0119gi g\u0142ównej"),
            ("4.", "Zestawienia obrotów i sald", "Sporz\u0105dzane miesi\u0119cznie dla kont ksi\u0119gi g\u0142ównej i ksi\u0105g pomocniczych"),
            ("5.", "Wykaz sk\u0142adników (inwentarz)", "Stosowany w przypadkach okre\u015blonych przepisami"),
        ],
        col_widths=[1, 4, 11]
    )

    doc.add_heading("\u00a7 9. System informatyczny", level=2)
    sf = G("d_sn") or "[nazwa programu ksi\u0119gowego]"
    sv_part = ", wersja: " + G("d_sv") if G("d_sv") else ""
    sp_part = ", producent: " + G("d_sp") if G("d_sp") else ""
    PJ("1. Ksi\u0119gi rachunkowe prowadzone s\u0105 przy u\u017cyciu systemu informatycznego: " + sf + sv_part + sp_part + ".")
    PJ("2. System informatyczny spe\u0142nia wymogi okre\u015blone w art. 10 ust. 1 pkt 3 lit. c oraz art. 13 ust. 2-6 Ustawy, w szczególno\u015bci zapewnia:")
    P("\u2022 trwa\u0142o\u015b\u0107 zapisu, niezmienno\u015b\u0107 wprowadzonych danych oraz zgodno\u015b\u0107 z dokumentami \u017aród\u0142owymi,", indent=False)
    P("\u2022 mo\u017cliwo\u015b\u0107 wydruku ksi\u0105g w postaci zestawie\u0144 ksi\u0119gowych za dowolny okres,", indent=False)
    P("\u2022 ochron\u0119 przed dost\u0119pem osób nieuprawnionych oraz przed utrat\u0105 danych,", indent=False)
    P("\u2022 prawid\u0142owo\u015b\u0107 zapisów ksi\u0119gowych oraz ich powi\u0105zanie z dowodami ksi\u0119gowymi.", indent=False)
    PJ("3. Szczegó\u0142owy opis systemu informatycznego, w tym wykaz programów wraz z pisemnym stwierdzeniem dopuszczenia ich do stosowania (art. 10 ust. 2 Ustawy), stanowi Za\u0142\u0105cznik nr 2 do niniejszej Polityki.")

    if G("d_ksef"):
        doc.add_heading("\u00a7 10. Krajowy System e-Faktur (KSeF)", level=2)
        PJ("1. Od dnia 1 lutego 2026 r. jednostka jako czynny podatnik VAT obowi\u0105zkowo uczestniczy w Krajowym Systemie e-Faktur (KSeF), zgodnie z ustaw\u0105 z dnia 16 czerwca 2023 r. o zmianie ustawy o podatku od towarów i us\u0142ug oraz niektórych innych ustaw (Dz.U. z 2023 r. poz. 1598).")
        PJ("2. Faktury sprzeda\u017cowe wystawiane s\u0105 wy\u0142\u0105cznie w formie ustrukturyzowanej (schemat FA(2)) i przesy\u0142ane do KSeF. Numer KSeF nadany przez system stanowi unikalny identyfikator faktury.")
        PJ("3. Faktury zakupowe otrzymywane s\u0105 za po\u015brednictwem KSeF. Weryfikacja poprawno\u015bci faktury (formalna i merytoryczna) nast\u0119puje po jej pobraniu z repozytorium KSeF.")

        ksef_mom = G("d_ksef_moment")
        if "wystawienia" in ksef_mom:
            PJ("4. Moment uj\u0119cia faktury w ksi\u0119gach rachunkowych: faktura ujmowana jest pod dat\u0105 wystawienia w KSeF, o ile data ta jest to\u017csama z dat\u0105 operacji gospodarczej. W przypadku rozbie\u017cno\u015bci \u2014 pod dat\u0105 operacji gospodarczej z adnotacj\u0105 o numerze KSeF i dacie wystawienia.")
        elif "operacji" in ksef_mom:
            PJ("4. Moment uj\u0119cia faktury w ksi\u0119gach rachunkowych: faktura ujmowana jest pod dat\u0105 operacji gospodarczej (data dostawy towaru lub wykonania us\u0142ugi), niezale\u017cnie od daty wystawienia w KSeF.")
        else:
            PJ("4. Moment uj\u0119cia faktury w ksi\u0119gach rachunkowych: faktura ujmowana jest pod dat\u0105 otrzymania (pobrania) z KSeF.")

        ksef_kor = G("d_ksef_korekty")
        if "Nota" in ksef_kor:
            PJ("5. Faktury koryguj\u0105ce: korekty danych formalnych realizowane s\u0105 w formie not koryguj\u0105cych przesy\u0142anych przez KSeF. Korekty warto\u015bciowe \u2014 w formie faktur koryguj\u0105cych z odniesieniem do numeru KSeF faktury pierwotnej.")
        elif "zbiorcza" in ksef_kor:
            PJ("5. Faktury koryguj\u0105ce: stosowane s\u0105 korekty zbiorcze za okresy rozliczeniowe, przesy\u0142ane przez KSeF z odniesieniem do okresu i kontrahenta.")
        else:
            PJ("5. Faktury koryguj\u0105ce: ka\u017cda korekta wystawiana jest jako odr\u0119bna faktura koryguj\u0105ca w KSeF z odniesieniem do numeru KSeF faktury pierwotnej.")

        ksef_sys = G("d_ksef_system")
        if "Zintegrowany" in ksef_sys:
            PJ("6. Integracja z systemem ksi\u0119gowym: system FK jednostki jest zintegrowany z KSeF poprzez API. Faktury sprzeda\u017cowe generowane s\u0105 automatycznie z systemu FK i przesy\u0142ane do KSeF. Faktury zakupowe importowane s\u0105 automatycznie z KSeF do systemu FK.")
        elif "automatyczny" in ksef_sys.lower():
            PJ("6. Integracja z systemem ksi\u0119gowym: faktury eksportowane/importowane s\u0105 w formacie XML (schemat FA(2)) mi\u0119dzy systemem FK a KSeF. Proces wymaga r\u0119cznego uruchomienia importu/eksportu w okresach miesi\u0119cznych.")
        else:
            PJ("6. Integracja z systemem ksi\u0119gowym: dane z faktur KSeF wprowadzane s\u0105 r\u0119cznie do systemu FK na podstawie podgl\u0105du faktury w repozytorium KSeF.")

        PJ("7. Numer KSeF przypisywany jest do ka\u017cdego zapisu ksi\u0119gowego dotycz\u0105cego faktury, umo\u017cliwiaj\u0105c pe\u0142n\u0105 identyfikowalno\u015b\u0107 na potrzeby JPK_VAT i JPK_CIT.")

    # ═══════ III. DOWODY KSIĘGOWE ═══════
    doc.add_heading("III. Dowody ksi\u0119gowe i ich obieg", level=1)

    doc.add_heading("\u00a7 11. Rodzaje dowodów ksi\u0119gowych", level=2)
    PJ("1. Podstaw\u0105 zapisów w ksi\u0119gach rachunkowych s\u0105 dowody ksi\u0119gowe stwierdzaj\u0105ce dokonanie operacji gospodarczej (art. 20 ust. 2 Ustawy).")
    PJ("2. Jednostka stosuje nast\u0119puj\u0105ce rodzaje dowodów ksi\u0119gowych:")
    add_table(
        ["Rodzaj dowodu", "Opis", "\u0179ród\u0142o"],
        [
            ("Zewn\u0119trzne obce", "Otrzymane od kontrahentów (faktury, rachunki, wezwania)", "Kontrahenci, KSeF"),
            ("Zewn\u0119trzne w\u0142asne", "Wystawiane na rzecz kontrahentów (faktury sprzeda\u017cy)", "Jednostka \u2192 KSeF"),
            ("Wewn\u0119trzne", "Dokumentuj\u0105ce operacje wewn\u0105trz jednostki (PK, PZ, WZ, LP)", "Jednostka"),
            ("Zbiorcze", "\u0141\u0105cz\u0105ce dowody jednorodne (zestawienia wp\u0142at, kompensaty)", "Jednostka"),
            ("Koryguj\u0105ce", "Korekty wcze\u015bniejszych dowodów (noty, faktury koryguj\u0105ce)", "Jednostka / KSeF"),
            ("Zast\u0119pcze", "Wystawiane w razie braku dowodu zewn\u0119trznego", "Jednostka"),
            ("Rozliczeniowe", "Stanowi\u0105ce podstaw\u0119 zapisów dekretowych (PK)", "Jednostka"),
        ],
        col_widths=[3.5, 8.5, 4]
    )

    doc.add_heading("\u00a7 12. Wymagania formalne dowodów", level=2)
    PJ("1. Dowód ksi\u0119gowy powinien zawiera\u0107 co najmniej elementy okre\u015blone w art. 21 ust. 1 Ustawy:")
    P("\u2022 okre\u015blenie rodzaju dowodu i jego numeru identyfikacyjnego,", indent=False)
    P("\u2022 okre\u015blenie stron operacji gospodarczej,", indent=False)
    P("\u2022 opis operacji oraz jej warto\u015b\u0107,", indent=False)
    P("\u2022 dat\u0119 dokonania operacji oraz dat\u0119 wystawienia dowodu,", indent=False)
    P("\u2022 podpis wystawcy i osoby, której powierzono sk\u0142adniki maj\u0105tku (je\u015bli wymagane),", indent=False)
    P("\u2022 stwierdzenie sprawdzenia i zakwalifikowania dowodu do uj\u0119cia w ksi\u0119gach (dekretacja).", indent=False)

    doc.add_heading("\u00a7 13. Obieg dokumentów", level=2)
    PJ("1. Obieg dokumentów w jednostce odbywa si\u0119 zgodnie z zasadami okre\u015blonymi w Instrukcji Obiegu Dokumentów, stanowi\u0105cej Za\u0142\u0105cznik nr 4 do niniejszej Polityki.")
    PJ("2. Ka\u017cdy dowód ksi\u0119gowy podlega kontroli merytorycznej, formalno-rachunkowej oraz dekretacji przed uj\u0119ciem w ksi\u0119gach rachunkowych.")

    # ═══════ IV. METODY WYCENY ═══════
    doc.add_heading("IV. Metody wyceny aktywów i pasywów", level=1)

    doc.add_heading("\u00a7 14. \u015arodki trwa\u0142e i warto\u015bci niematerialne i prawne", level=2)
    PJ("1. \u015arodki trwa\u0142e oraz warto\u015bci niematerialne i prawne o warto\u015bci pocz\u0105tkowej przekraczaj\u0105cej " + thr + " PLN ujmowane s\u0105 w ewidencji \u015brodków trwa\u0142ych i amortyzowane zgodnie z planem amortyzacji.")
    dm = {"Metoda liniowa": "metod\u0105 liniow\u0105", "Metoda degresywna": "metod\u0105 degresywn\u0105 (wspó\u0142czynnik 2,0)", "Jednorazowa": "jednorazowo, do limitu ustawowego"}
    PJ("2. \u015arodki trwa\u0142e amortyzowane s\u0105 " + dm.get(G("d_dep"), G("d_dep")) + ", zgodnie ze stawkami wynikaj\u0105cymi z wykazu stawek amortyzacyjnych stanowi\u0105cego za\u0142\u0105cznik do ustawy o podatku dochodowym od osób prawnych.")
    PJ("3. Sk\u0142adniki maj\u0105tku o warto\u015bci pocz\u0105tkowej nieprzekraczaj\u0105cej " + thr + " PLN mog\u0105 by\u0107 jednorazowo odpisywane w koszty w miesi\u0105cu oddania do u\u017cytkowania, bez ujmowania w ewidencji \u015brodków trwa\u0142ych.")
    PJ("4. Warto\u015bci niematerialne i prawne amortyzowane s\u0105 metod\u0105 liniow\u0105 przez okres ekonomicznej u\u017cyteczno\u015bci:")
    add_table(
        ["Rodzaj WNiP", "Okres amortyzacji"],
        [
            ("Licencje na oprogramowanie komputerowe", "minimum 24 miesi\u0105ce"),
            ("Prawa autorskie i pokrewne", "minimum 24 miesi\u0105ce"),
            ("Koszty zako\u0144czonych prac rozwojowych", "maksymalnie 5 lat"),
            ("Warto\u015b\u0107 firmy", "maksymalnie 5 lat (art. 44b ust. 10 Ustawy)"),
            ("Pozosta\u0142e WNiP", "minimum 60 miesi\u0119cy"),
        ],
        col_widths=[8, 8]
    )
    PJ("5. Warto\u015b\u0107 pocz\u0105tkowa \u015brodków trwa\u0142ych obejmuje cen\u0119 nabycia lub koszt wytworzenia powi\u0119kszone o koszty bezpo\u015brednio zwi\u0105zane z zakupem i przystosowaniem do u\u017cywania (monta\u017c, transport, op\u0142aty notarialne, odsetki od kredytów do momentu oddania do u\u017cywania).")
    PJ("6. \u015arodki trwa\u0142e w budowie wycenia si\u0119 w wysoko\u015bci ogó\u0142u kosztów pozostaj\u0105cych w bezpo\u015brednim zwi\u0105zku z ich budow\u0105, monta\u017cem lub przystosowaniem, pomniejszonych o odpisy z tytu\u0142u trwa\u0142ej utraty warto\u015bci.")
    PJ("7. Trwa\u0142\u0105 utrat\u0119 warto\u015bci \u015brodków trwa\u0142ych ocenia si\u0119 na ka\u017cdy dzie\u0144 bilansowy zgodnie z KSR 4. W razie stwierdzenia takiej utraty dokonuje si\u0119 odpisu aktualizuj\u0105cego.")

    doc.add_heading("\u00a7 15. Zapasy", level=2)
    ivm = {"Cena nabycia": "cen nabycia", "Koszt wytworzenia": "kosztu wytworzenia", "Cena rynkowa": "warto\u015bci rynkowej"}
    PJ("1. Zapasy (rzeczowe aktywa obrotowe) wyceniane s\u0105 wed\u0142ug " + ivm.get(G("d_iv"), G("d_iv")) + ", nie wy\u017cszych od cen sprzeda\u017cy netto na dzie\u0144 bilansowy.")
    PJ("2. Cena nabycia obejmuje cen\u0119 zakupu powi\u0119kszon\u0105 o koszty bezpo\u015brednio zwi\u0105zane z zakupem (c\u0142a, transport, ubezpieczenie, op\u0142aty publicznoprawne) oraz pomniejszon\u0105 o rabaty, opusty i inne podobne zmniejszenia.")
    idm = {"FIFO": "FIFO (pierwsze wesz\u0142o \u2014 pierwsze wysz\u0142o)", "LIFO": "LIFO (ostatnie wesz\u0142o \u2014 pierwsze wysz\u0142o)", "Srednia wazona": "\u015bredniej wa\u017conej", "Szczegolowa identyfikacja": "szczegó\u0142owej identyfikacji"}
    PJ("3. Rozchód zapasów wyceniany jest metod\u0105 " + idm.get(G("d_id"), G("d_id")) + ", zgodnie z art. 34 ust. 4 Ustawy. Wybran\u0105 metod\u0119 stosuje si\u0119 konsekwentnie w stosunku do zapasów o podobnym charakterze i przeznaczeniu.")
    PJ("4. Zapasy o obni\u017conej przydatno\u015bci gospodarczej lub utraconej warto\u015bci handlowej obejmowane s\u0105 odpisami aktualizuj\u0105cymi do warto\u015bci netto mo\u017cliwej do uzyskania.")
    PJ("5. Inwentaryzacja zapasów przeprowadzana jest zgodnie z zasadami okre\u015blonymi w Rozdziale VII niniejszej Polityki.")

    doc.add_heading("\u00a7 16. Nale\u017cno\u015bci", level=2)
    PJ("1. Nale\u017cno\u015bci wycenia si\u0119 w kwocie wymaganej zap\u0142aty, z zachowaniem zasady ostro\u017cno\u015bci (art. 28 ust. 1 pkt 7 Ustawy), tj. po pomniejszeniu o odpisy aktualizuj\u0105ce.")
    PJ("2. Odpisy aktualizuj\u0105ce warto\u015b\u0107 nale\u017cno\u015bci tworzy si\u0119 w odniesieniu do:")
    P("\u2022 nale\u017cno\u015bci od d\u0142u\u017cników postawionych w stan likwidacji lub upad\u0142o\u015bci,", indent=False)
    P("\u2022 nale\u017cno\u015bci kwestionowanych przez d\u0142u\u017cników lub z których zap\u0142at\u0105 d\u0142u\u017cnik zalega,", indent=False)
    P("\u2022 nale\u017cno\u015bci przeterminowanych powy\u017cej 180 dni \u2014 w wysoko\u015bci 50%,", indent=False)
    P("\u2022 nale\u017cno\u015bci przeterminowanych powy\u017cej 360 dni \u2014 w wysoko\u015bci 100%.", indent=False)
    PJ("3. Nale\u017cno\u015bci w walutach obcych wycenia si\u0119 zgodnie z zasadami okre\u015blonymi w Rozdziale VI niniejszej Polityki.")

    doc.add_heading("\u00a7 17. Inwestycje", level=2)
    PJ("1. Inwestycje krótkoterminowe (papiery warto\u015bciowe, udzia\u0142y, lokaty) wycenia si\u0119 wed\u0142ug ceny nabycia lub warto\u015bci rynkowej, zale\u017cnie od tego, która z nich jest ni\u017csza (art. 28 ust. 1 pkt 5 Ustawy).")
    PJ("2. Inwestycje d\u0142ugoterminowe (udzia\u0142y i akcje w jednostkach powi\u0105zanych i pozosta\u0142ych) wycenia si\u0119 wed\u0142ug ceny nabycia pomniejszonej o odpisy z tytu\u0142u trwa\u0142ej utraty warto\u015bci.")
    PJ("3. Nieruchomo\u015bci inwestycyjne wycenia si\u0119 wed\u0142ug zasad stosowanych do \u015brodków trwa\u0142ych.")

    doc.add_heading("\u00a7 18. Zobowi\u0105zania i rezerwy", level=2)
    PJ("1. Zobowi\u0105zania wycenia si\u0119 w kwocie wymagaj\u0105cej zap\u0142aty (art. 28 ust. 1 pkt 8 Ustawy).")
    PJ("2. Zobowi\u0105zania finansowe (kredyty, po\u017cyczki) wycenia si\u0119 wed\u0142ug skorygowanej ceny nabycia z zastosowaniem efektywnej stopy procentowej.")
    PJ("3. Rezerwy tworzy si\u0119 na pewne lub o du\u017cym stopniu prawdopodobie\u0144stwa przysz\u0142e zobowi\u0105zania (art. 35d Ustawy). Obejmuj\u0105 w szczególno\u015bci:")
    P("\u2022 rezerwy na \u015bwiadczenia pracownicze (odprawy emerytalne, nagrody jubileuszowe, niewykorzystane urlopy),", indent=False)
    P("\u2022 rezerwy na naprawy gwarancyjne i r\u0119kojmi\u0119,", indent=False)
    P("\u2022 rezerwy na sprawy sporne i post\u0119powania s\u0105dowe,", indent=False)
    P("\u2022 rezerwy na restrukturyzacj\u0119.", indent=False)
    PJ("4. Bierne rozliczenia mi\u0119dzyokresowe kosztów dokonywane s\u0105 zgodnie z art. 39 ust. 2 Ustawy w wysoko\u015bci prawdopodobnych zobowi\u0105za\u0144 przypadaj\u0105cych na bie\u017c\u0105cy okres sprawozdawczy.")

    doc.add_heading("\u00a7 19. Kapita\u0142y w\u0142asne", level=2)
    PJ("1. Kapita\u0142 zak\u0142adowy wykazuje si\u0119 w wysoko\u015bci okre\u015blonej w umowie spó\u0142ki, ujawnionej w Krajowym Rejestrze S\u0105dowym.")
    PJ("2. Kapita\u0142 zapasowy tworzony jest z odpisów z zysku oraz innych \u017aróde\u0142 okre\u015blonych w przepisach prawa i umowie spó\u0142ki.")
    PJ("3. Niepodzielone wyniki finansowe z lat ubieg\u0142ych wykazuje si\u0119 w odr\u0119bnej pozycji bilansu.")

    # ═══════ V. KOSZTY I RZIS ═══════
    doc.add_heading("V. Ewidencja kosztów i wariant Rachunku Zysków i Strat", level=1)

    doc.add_heading("\u00a7 20. Model ewidencji kosztów", level=2)
    cmm = {
        "Tylko Zespol 4 (uklad rodzajowy)": "wy\u0142\u0105cznie w Zespole 4 (uk\u0142ad rodzajowy)",
        "Tylko Zespol 5 (uklad kalkulacyjny)": "wy\u0142\u0105cznie w Zespole 5 (uk\u0142ad funkcjonalno-kalkulacyjny)",
        "Zespol 4 + 5 (oba uklady)": "równolegle w Zespole 4 i Zespole 5"
    }
    PJ("1. Jednostka prowadzi ewidencj\u0119 kosztów " + cmm.get(G("d_cm"), G("d_cm")) + ".")

    if "Zespol 4" in G("d_cm"):
        PJ("2. Koszty ujmowane s\u0105 w uk\u0142adzie rodzajowym, zgodnie z tre\u015bci\u0105 ekonomiczn\u0105 operacji gospodarczej, w nast\u0119puj\u0105cych kategoriach:")
        add_table(
            ["Konto", "Nazwa", "Opis"],
            [
                ("400", "Amortyzacja", "Odpisy amortyzacyjne ST i WNiP"),
                ("401", "Zu\u017cycie materia\u0142ów i energii", "Materia\u0142y, energia, paliwo"),
                ("402", "Us\u0142ugi obce", "Us\u0142ugi transportowe, remontowe, doradcze, najem"),
                ("403", "Podatki i op\u0142aty", "Podatki niezaliczone do CIT, op\u0142aty urz\u0119dowe"),
                ("404", "Wynagrodzenia", "Wynagrodzenia brutto pracowników i zleceniobiorców"),
                ("405", "Ubezpieczenia spo\u0142eczne i \u015bwiadczenia", "ZUS, BHP, \u015bwiadczenia socjalne"),
                ("409", "Pozosta\u0142e koszty rodzajowe", "Koszty niemieszcz\u0105ce si\u0119 w innych kategoriach"),
            ],
            col_widths=[2, 5, 9]
        )
    if "Zespol 5" in G("d_cm") or "4 + 5" in G("d_cm"):
        PJ("3. Koszty ujmowane s\u0105 równie\u017c w uk\u0142adzie funkcjonalno-kalkulacyjnym (Zespó\u0142 5), w podziale na:")
        P("\u2022 koszty produkcji podstawowej (501),", indent=False)
        P("\u2022 koszty wydzia\u0142ów (520),", indent=False)
        P("\u2022 koszty sprzeda\u017cy (527),", indent=False)
        P("\u2022 koszty ogólnego zarz\u0105du (550).", indent=False)

    doc.add_heading("\u00a7 21. Wariant Rachunku Zysków i Strat", level=2)
    plbl = "porównawczym" if "porownawczy" in G("d_pl") else "kalkulacyjnym"
    atn = "4" if G("d_micro") else "5" if G("d_small") else "1"
    PJ("1. Rachunek Zysków i Strat sporz\u0105dzany jest w wariancie " + plbl + ", zgodnie z Za\u0142\u0105cznikiem nr " + atn + " do Ustawy.")
    if "porownawczy" in G("d_pl"):
        PJ("2. W wariancie porównawczym koszty i przychody prezentowane s\u0105 w uk\u0142adzie rodzajowym, ze zmian\u0105 stanu produktów jako oddzieln\u0105 pozycj\u0105 koryguj\u0105c\u0105.")
    else:
        PJ("2. W wariancie kalkulacyjnym koszty przypisywane s\u0105 do funkcji (produkcja, sprzeda\u017c, zarz\u0105d), a przychody netto ze sprzeda\u017cy konfrontowane s\u0105 z kosztem wytworzenia sprzedanych produktów.")

    # ═══════ VI. WALUTY OBCE ═══════
    doc.add_heading("VI. Operacje gospodarcze w walutach obcych", level=1)

    doc.add_heading("\u00a7 22. Kurs walut", level=2)
    fxm = {
        "Kurs sredni NBP": "\u015brednim og\u0142aszanym przez Narodowy Bank Polski z ostatniego dnia roboczego poprzedzaj\u0105cego dzie\u0144 operacji",
        "Kurs kupna banku": "kupna banku, z którego us\u0142ug korzysta jednostka",
        "Kurs sprzedazy banku": "sprzeda\u017cy banku, z którego us\u0142ug korzysta jednostka"
    }
    PJ("1. Operacje gospodarcze wyra\u017cone w walutach obcych ujmuje si\u0119 w ksi\u0119gach rachunkowych w walucie polskiej, przeliczone po kursie " + fxm.get(G("d_fxs"), G("d_fxs")) + " (art. 30 ust. 2 Ustawy).")
    PJ("2. Na dzie\u0144 bilansowy aktywa i pasywa wyra\u017cone w walutach obcych wycenia si\u0119 po kursie \u015brednim og\u0142oszonym przez NBP na ten dzie\u0144, zgodnie z art. 30 ust. 1 Ustawy.")
    PJ("3. Ró\u017cnice kursowe (zrealizowane i niezrealizowane) odnosi si\u0119 odpowiednio na przychody finansowe (dodatnie) lub koszty finansowe (ujemne).")

    if G("d_hfx"):
        doc.add_heading("\u00a7 23. Rachunki walutowe", level=2)
        cdm = {"FIFO": "FIFO", "LIFO": "LIFO", "Srednia wazona": "\u015bredniej wa\u017conej"}
        cur = G("d_cur")
        cur_list = ", ".join(cur) if isinstance(cur, list) and cur else "EUR, USD"
        PJ("1. Jednostka prowadzi rachunki walutowe w nast\u0119puj\u0105cych walutach: " + cur_list + ".")
        PJ("2. Rozchód \u015brodków pieni\u0119\u017cnych z rachunków walutowych wyceniany jest metod\u0105 " + cdm.get(G("d_fxd"), G("d_fxd")) + ".")
        PJ("3. Stan \u015brodków pieni\u0119\u017cnych na rachunkach walutowych na dzie\u0144 bilansowy wycenia si\u0119 po kursie \u015brednim NBP z dnia bilansowego.")

    # ═══════ VII. INWENTARYZACJA ═══════
    doc.add_heading("VII. Inwentaryzacja", level=1)

    doc.add_heading("\u00a7 24. Cel i zakres inwentaryzacji", level=2)
    PJ("1. Inwentaryzacja przeprowadzana jest zgodnie z art. 26 i 27 Ustawy oraz Krajowym Standardem Rachunkowo\u015bci nr 1 dotycz\u0105cym inwentaryzacji.")
    PJ("2. Celem inwentaryzacji jest ustalenie rzeczywistego stanu sk\u0142adników aktywów i pasywów oraz porównanie go ze stanem ksi\u0119gowym.")

    doc.add_heading("\u00a7 25. Metody i terminy inwentaryzacji", level=2)
    add_table(
        ["Sk\u0142adnik", "Metoda", "Termin", "Cz\u0119stotliwo\u015b\u0107"],
        [
            ("\u015arodki pieni\u0119\u017cne w kasie", "Spis z natury", "31.12.", "Roczna"),
            ("Druki \u015bcis\u0142ego zarachowania", "Spis z natury", "31.12.", "Roczna"),
            ("Zapasy (towary, materia\u0142y)", "Spis z natury", "Q4 / 15.10\u201415.01", "Roczna"),
            ("\u015arodki trwa\u0142e", "Spis z natury", "Q4", "Co 4 lata"),
            ("\u015arodki pieni\u0119\u017cne na rachunkach", "Potwierdzenie salda", "31.12.", "Roczna"),
            ("Nale\u017cno\u015bci i zobowi\u0105zania", "Potwierdzenie salda", "Q4", "Roczna"),
            ("Pozosta\u0142e aktywa i pasywa", "Weryfikacja", "31.12.", "Roczna"),
        ],
        col_widths=[4.5, 3.5, 3.5, 3]
    )

    doc.add_heading("\u00a7 26. Komisja inwentaryzacyjna", level=2)
    PJ("1. Inwentaryzacj\u0119 przeprowadza komisja inwentaryzacyjna powo\u0142ana zarz\u0105dzeniem kierownika jednostki.")
    PJ("2. Komisja sporz\u0105dza protokó\u0142 inwentaryzacyjny zawieraj\u0105cy zestawienie ró\u017cnic inwentaryzacyjnych (nadwy\u017cek i niedoborów) wraz z wyja\u015bnieniem ich przyczyn.")
    PJ("3. Rozliczenie ró\u017cnic inwentaryzacyjnych nast\u0119puje uchwa\u0142\u0105 lub decyzj\u0105 kierownika jednostki na podstawie protoko\u0142u komisji.")

    # ═══════ VIII. OCHRONA DANYCH ═══════
    doc.add_heading("VIII. System ochrony danych i ich zbiorów", level=1)

    doc.add_heading("\u00a7 27. Zasady ochrony danych", level=2)
    dpm = {
        "Elektroniczna i fizyczna": "elektronicznej i fizycznej",
        "Wylacznie elektroniczna": "wy\u0142\u0105cznie elektronicznej",
        "Wylacznie fizyczna": "wy\u0142\u0105cznie fizycznej"
    }
    PJ("1. Ochrona danych rachunkowych realizowana jest w formie " + dpm.get(G("d_dp"), G("d_dp")) + ", zgodnie z art. 71-72 Ustawy.")
    PJ("2. Jednostka zapewnia trwa\u0142o\u015b\u0107 zapisów ksi\u0119gowych, ich nienaruszalno\u015b\u0107, ochron\u0119 przed dost\u0119pem osób nieuprawnionych oraz przed utrat\u0105.")

    doc.add_heading("\u00a7 28. Kopie zapasowe", level=2)
    bkm = {"Codziennie": "codziennej", "Co tydzien": "tygodniowej", "Co miesiac": "miesi\u0119cznej"}
    PJ("1. Kopie zapasowe danych ksi\u0119gowych sporz\u0105dzane s\u0105 z cz\u0119stotliwo\u015bci\u0105 " + bkm.get(G("d_bk"), G("d_bk")) + ", na no\u015bnikach zapewniaj\u0105cych trwa\u0142o\u015b\u0107 zapisu.")
    PJ("2. Kopie zapasowe przechowywane s\u0105 w lokalizacji zewn\u0119trznej (poza siedzib\u0105 jednostki) lub na chronionym serwerze w chmurze, zapewniaj\u0105c odzyskanie danych w przypadku awarii systemu.")
    PJ("3. Przeprowadzane s\u0105 okresowe testy odtwarzania danych z kopii zapasowych, nie rzadziej ni\u017c raz w roku.")

    doc.add_heading("\u00a7 29. Kontrola dost\u0119pu", level=2)
    if G("d_ac"):
        PJ("1. Dost\u0119p do systemu ksi\u0119gowego zabezpieczony jest indywidualnymi loginami i has\u0142ami u\u017cytkowników.")
        PJ("2. Has\u0142a zmieniane s\u0105 nie rzadziej ni\u017c co 90 dni i musz\u0105 spe\u0142nia\u0107 wymagania z\u0142o\u017cono\u015bci (minimum 8 znaków, w tym wielkie i ma\u0142e litery, cyfry oraz znaki specjalne).")
        PJ("3. Uprawnienia u\u017cytkowników nadawane s\u0105 na zasadzie minimum koniecznego \u2014 wy\u0142\u0105cznie do funkcji i danych niezb\u0119dnych do wykonywania obowi\u0105zków.")
    else:
        PJ("1. Jednostka zapewnia odpowiedni poziom kontroli dost\u0119pu do danych rachunkowych zgodnie z opracowan\u0105 procedur\u0105 wewn\u0119trzn\u0105.")

    doc.add_heading("\u00a7 30. Archiwizacja", level=2)
    PJ("1. Dokumentacj\u0119 ksi\u0119gow\u0105 przechowuje si\u0119 przez okres " + str(G("d_ay")) + " lat, liczony od pocz\u0105tku roku nast\u0119puj\u0105cego po roku obrotowym, którego dane zbiory dotycz\u0105 (art. 74 Ustawy).")
    PJ("2. Sprawozdania finansowe oraz roczne deklaracje podatkowe przechowuje si\u0119 trwale.")
    PJ("3. Dokumenty pracownicze (akta osobowe, listy p\u0142ac) przechowywane s\u0105 przez okres wynikaj\u0105cy z przepisów odr\u0119bnych \u2014 10 lub 50 lat.")

    if G("d_ksef"):
        doc.add_heading("\u00a7 31. Archiwizacja faktur KSeF", level=2)
        ksef_arch = G("d_ksef_archiwum")
        if "Wylacznie" in ksef_arch:
            PJ("1. Faktury ustrukturyzowane przechowywane s\u0105 wy\u0142\u0105cznie w repozytorium KSeF prowadzonym przez Ministerstwo Finansów. Repozytorium KSeF spe\u0142nia wymogi art. 73 i 74 Ustawy w zakresie trwa\u0142o\u015bci zapisu i ochrony przed modyfikacj\u0105.")
        elif "systemie FK" in ksef_arch and "nosnikach" not in ksef_arch:
            PJ("1. Faktury ustrukturyzowane przechowywane s\u0105 równolegle w repozytorium KSeF oraz w systemie FK jednostki. System FK przechowuje kopie faktur w formacie XML (schemat FA(2)) wraz z numerem KSeF.")
        else:
            PJ("1. Faktury ustrukturyzowane przechowywane s\u0105 w trzech lokalizacjach: repozytorium KSeF, system FK jednostki oraz na no\u015bnikach zapasowych \u2014 co zapewnia pe\u0142n\u0105 redundancj\u0119 i ci\u0105g\u0142o\u015b\u0107 dost\u0119pu do dokumentów.")
        PJ("2. Okres przechowywania faktur KSeF: zgodnie z art. 74 Ustawy oraz art. 112 ustawy o VAT \u2014 co najmniej 5 lat, liczone od ko\u0144ca roku kalendarzowego, w którym up\u0142yn\u0105\u0142 termin p\u0142atno\u015bci podatku.")

    PJ("4. Osoba odpowiedzialna za system ochrony danych w jednostce: " + (G("d_rp") or "[imi\u0119 i nazwisko]") + ".")
    PJ("5. Szczegó\u0142owy opis systemu ochrony danych, w tym procedury awaryjne, instrukcja post\u0119powania w przypadku incydentu, stanowi Za\u0142\u0105cznik nr 3 do niniejszej Polityki.")

    # ═══════ IX. JPK I RAPORTOWANIE ═══════
    doc.add_heading("IX. Raportowanie elektroniczne (JPK, KSeF)", level=1)

    doc.add_heading("\u00a7 32. Jednolity Plik Kontrolny", level=2)
    PJ("1. Jednostka udost\u0119pnia organom podatkowym Jednolite Pliki Kontrolne (JPK) zgodnie z wymaganiami ustawy z dnia 29 sierpnia 1997 r. \u2014 Ordynacja podatkowa oraz przepisów wykonawczych:")
    add_table(
        ["Struktura JPK", "Zakres", "Termin"],
        [
            ("JPK_VAT (JPK_V7M/V7K)", "Ewidencja VAT \u2014 rejestr sprzeda\u017cy i zakupu", "Miesi\u0119cznie / kwartalnie"),
            ("JPK_KR", "Ksi\u0119gi rachunkowe", "Na \u017c\u0105danie organu"),
            ("JPK_MAG", "Magazyn (przyj\u0119cia, wydania)", "Na \u017c\u0105danie organu"),
            ("JPK_WB", "Wyci\u0105gi bankowe", "Na \u017c\u0105danie organu"),
            ("JPK_FA", "Faktury VAT", "Na \u017c\u0105danie organu (od 2026 \u2014 przez KSeF)"),
            ("JPK_CIT (od 2026)", "Ksi\u0119gi rachunkowe w nowej strukturze CIT", "Rocznie"),
        ],
        col_widths=[4, 8, 4]
    )
    PJ("2. Plan kont jednostki jest skonstruowany w sposób umo\u017cliwiaj\u0105cy automatyczne mapowanie zapisów ksi\u0119gowych na pozycje JPK_CIT od roku obrotowego 2026.")

    # ═══════ X. ZASADY DODATKOWE ═══════
    doc.add_heading("X. Zasady dodatkowe i polityki szczególne", level=1)

    doc.add_heading("\u00a7 33. Ujmowanie przychodów", level=2)
    if "memorialowa" in G("d_rev"):
        PJ("1. Przychody ujmowane s\u0105 zgodnie z zasad\u0105 memoria\u0142ow\u0105 (art. 6 ust. 1 Ustawy), tj. w okresie, którego dotycz\u0105, niezale\u017cnie od terminu wp\u0142ywu \u015brodków pieni\u0119\u017cnych.")
        PJ("2. Przychody ze sprzeda\u017cy produktów, towarów i us\u0142ug ujmuje si\u0119 w momencie:")
        P("\u2022 dostarczenia produktu lub towaru kupuj\u0105cemu,", indent=False)
        P("\u2022 wykonania us\u0142ugi lub jej etapu zgodnie z umow\u0105,", indent=False)
        P("\u2022 gdy kwot\u0119 przychodu mo\u017cna w sposób wiarygodny okre\u015bli\u0107,", indent=False)
        P("\u2022 gdy istnieje prawdopodobie\u0144stwo uzyskania korzy\u015bci ekonomicznych.", indent=False)
    else:
        PJ("1. Przychody ujmowane s\u0105 zgodnie z zasad\u0105 kasow\u0105, tj. w momencie otrzymania zap\u0142aty.")

    doc.add_heading("\u00a7 34. Klasyfikacja umów leasingu", level=2)
    if "bilansow" in G("d_ls"):
        PJ("1. Umowy leasingu klasyfikowane s\u0105 zgodnie z przepisami bilansowymi (art. 3 ust. 4-6 Ustawy).")
        PJ("2. Leasing finansowy ujmowany jest w aktywach jednostki korzystaj\u0105cej, z jednoczesnym uj\u0119ciem zobowi\u0105zania wobec finansuj\u0105cego.")
        PJ("3. Leasing operacyjny ujmowany jest w kosztach okresu w wysoko\u015bci rat leasingowych przypadaj\u0105cych na ten okres.")
    else:
        PJ("1. Umowy leasingu klasyfikowane s\u0105 zgodnie z przepisami podatkowymi (art. 17a-17l ustawy o CIT).")
        PJ("2. Przedmiot leasingu ujmowany jest w aktywach finansuj\u0105cego. U korzystaj\u0105cego raty leasingowe ujmowane s\u0105 jako koszt okresu.")

    if G("d_prov"):
        doc.add_heading("\u00a7 35. Rezerwy i rozliczenia mi\u0119dzyokresowe", level=2)
        PJ("1. Jednostka tworzy rezerwy na znane ryzyko, gro\u017c\u0105ce straty oraz skutki innych zdarze\u0144, zgodnie z art. 35d Ustawy.")
        PJ("2. Bierne rozliczenia mi\u0119dzyokresowe kosztów dokonywane s\u0105 zgodnie z art. 39 Ustawy w wysoko\u015bci prawdopodobnych zobowi\u0105za\u0144 przypadaj\u0105cych na bie\u017c\u0105cy okres sprawozdawczy.")
        PJ("3. Czynne rozliczenia mi\u0119dzyokresowe kosztów dotycz\u0105 wydatków poniesionych w bie\u017c\u0105cym okresie, a obci\u0105\u017caj\u0105cych koszty okresów przysz\u0142ych.")

    if G("d_dt"):
        doc.add_heading("\u00a7 36. Podatek odroczony", level=2)
        if G("d_small") or G("d_micro"):
            PJ("1. Jednostka, jako jednostka ma\u0142a/mikro, korzysta z uproszczenia polegaj\u0105cego na zaniechaniu ustalania aktywów i rezerw z tytu\u0142u odroczonego podatku dochodowego, zgodnie z art. 37 ust. 10 Ustawy.")
        else:
            PJ("1. Jednostka ustala aktywa i rezerwy z tytu\u0142u odroczonego podatku dochodowego, zgodnie z art. 37 Ustawy.")
            PJ("2. Aktywa z tytu\u0142u odroczonego podatku ustala si\u0119 w wysoko\u015bci kwoty przewidzianej do odliczenia od podatku dochodowego w przysz\u0142o\u015bci.")
            PJ("3. Rezerwy z tytu\u0142u odroczonego podatku tworzy si\u0119 w wysoko\u015bci kwoty podatku do zap\u0142acenia w przysz\u0142o\u015bci.")

    doc.add_heading("\u00a7 37. Rachunek przep\u0142ywów pieni\u0119\u017cnych", level=2)
    if G("d_small") or G("d_micro"):
        PJ("1. Jednostka, jako jednostka ma\u0142a/mikro, jest zwolniona z obowi\u0105zku sporz\u0105dzania rachunku przep\u0142ywów pieni\u0119\u017cnych.")
    else:
        cf = "po\u015bredni\u0105" if "posrednia" in G("d_cf") else "bezpo\u015bredni\u0105"
        PJ("1. Rachunek przep\u0142ywów pieni\u0119\u017cnych sporz\u0105dzany jest metod\u0105 " + cf + ", zgodnie z Za\u0142\u0105cznikiem nr 1 do Ustawy oraz KSR 1.")

    # ═══════ XI. POSTANOWIENIA KOŃCOWE ═══════
    doc.add_heading("XI. Postanowienia ko\u0144cowe", level=1)

    doc.add_heading("\u00a7 38. Wej\u015bcie w \u017cycie", level=2)
    PJ("1. Niniejsza Polityka Rachunkowo\u015bci wchodzi w \u017cycie z dniem " + eds + " i obowi\u0105zuje do czasu jej zmiany.")
    PJ("2. Polityka obowi\u0105zuje od roku obrotowego rozpoczynaj\u0105cego si\u0119 w dniu wej\u015bcia w \u017cycie.")

    doc.add_heading("\u00a7 39. Zmiany Polityki", level=2)
    PJ("1. Wszelkie zmiany niniejszej Polityki wymagaj\u0105 formy pisemnej i zatwierdzenia przez " + (G("d_ab") or "kierownika jednostki") + ".")
    PJ("2. Zmiany Polityki dokonuje si\u0119 z zachowaniem zasady ci\u0105g\u0142o\u015bci (art. 5 ust. 1 Ustawy). Skutki zmian odnosi si\u0119 na kapita\u0142 w\u0142asny, je\u015bli wymagaj\u0105 tego okoliczno\u015bci (KSR 7).")
    PJ("3. Zmiany wprowadza si\u0119 w drodze aneksu do niniejszej Polityki, ze wskazaniem daty wprowadzenia i przyczyny zmiany.")

    doc.add_heading("\u00a7 40. Odpowiedzialno\u015b\u0107", level=2)
    PJ("1. Za przestrzeganie zasad (polityki) rachunkowo\u015bci odpowiada kierownik jednostki, zgodnie z art. 4 ust. 5 Ustawy.")
    PJ("2. Pracownicy dzia\u0142u ksi\u0119gowo\u015bci oraz inne osoby uczestnicz\u0105ce w procesach ksi\u0119gowych zobowi\u0105zane s\u0105 do przestrzegania niniejszej Polityki.")

    # ═══════ ZAŁĄCZNIKI ═══════
    doc.add_heading("Za\u0142\u0105czniki", level=1)
    PJ("Integraln\u0105 cz\u0119\u015bci\u0105 niniejszej Polityki Rachunkowo\u015bci s\u0105 nast\u0119puj\u0105ce za\u0142\u0105czniki:")
    add_table(
        ["Nr", "Nazwa za\u0142\u0105cznika", "Opis"],
        [
            ("1", "Zak\u0142adowy Plan Kont", "Wykaz kont ksi\u0119gi g\u0142ównej i ksi\u0105g pomocniczych wraz z opisem"),
            ("2", "Opis systemu informatycznego", "Wykaz programów ksi\u0119gowych wraz z dokumentacj\u0105"),
            ("3", "System ochrony danych", "Procedury ochrony, kopii zapasowych i odzyskiwania danych"),
            ("4", "Instrukcja obiegu dokumentów", "Zasady obiegu dowodów ksi\u0119gowych w jednostce"),
            ("5", "Instrukcja inwentaryzacyjna", "Procedury i terminy inwentaryzacji"),
            ("6", "Wykaz osób upowa\u017cnionych", "Pracownicy uprawnieni do zatwierdzania dowodów ksi\u0119gowych"),
        ],
        col_widths=[1.5, 5, 9.5]
    )

    # Podpisy
    doc.add_paragraph()
    doc.add_paragraph()
    PJ("Sporz\u0105dzi\u0142(a):                                              Zatwierdzi\u0142(a):", indent=False)
    doc.add_paragraph()
    PJ("____________________________                __________________________________", indent=False)
    PJ("            (G\u0142ówny Ksi\u0119gowy)                                  " + (G("d_ab") or "(Kierownik jednostki)"), indent=False)
    doc.add_paragraph()
    PC("Data: " + ads, 10, c="666666")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════
# WIZARD STEPS
# ══════════════════════════════════════════════════════

def step_0():
    st.subheader("Krok 1: Sprawdz dane jednostki")
    st.info("Dane jednostki wypelnij w **panelu bocznym po lewej**. Mozesz pobrac je z KRS.")
    krs = st.session_state.get("krs_data", {})
    if krs.get("nazwa"):
        st.success(f"Dane z KRS: **{krs['nazwa']}**")
    st.write(f"**Nazwa:** {G('d_name') or '-'}")
    st.write(f"**NIP:** {G('d_nip') or '-'} | **KRS:** {G('d_krs') or '-'} | **REGON:** {G('d_regon') or '-'}")
    st.write(f"**Adres:** {G('d_addr') or '-'}")


def step_1():
    st.subheader("Krok 2: Ksiegi rachunkowe i Plan Kont")

    st.session_state.d_zpk = st.radio("Zakladowy Plan Kont",
        ["Wzorcowy plan kont", "Wygeneruj plan kont na podstawie parametrow"], key="wzpk")

    if "Wygeneruj" in st.session_state.d_zpk:
        st.markdown("---")
        st.markdown("### Generator Zakladowego Planu Kont (ZPK) 2026")
        st.caption("Odpowiedz na pytania - system wygeneruje ZPK z uwzglednieniem JPK_CIT i KSeF.")

        c1, c2 = st.columns(2)
        with c1:
            zpk_branza = st.selectbox("Branza", ["Uslugi", "Handel", "Produkcja", "Hybryda"], key="zpk_br")
            zpk_cit = st.selectbox("Typ CIT", ["Klasyczny", "Estonski"], key="zpk_cit")
            zpk_rzis = st.selectbox("Wariant RZiS", ["Porownawczy", "Kalkulacyjny"], key="zpk_rzis")
        with c2:
            zpk_skala = st.selectbox("Skala podatnika", ["Maly", "Duzy"], key="zpk_sk")
            zpk_aut = st.selectbox("Analityka KUP/NKUP (samochody, limity)", ["Tak", "Nie"], key="zpk_aut")
            zpk_tp = st.selectbox("Podmioty powiazane (TP)", ["Nie", "Tak"], key="zpk_tp")

        if st.button("Generuj Plan Kont", use_container_width=True, type="primary", key="gen_zpk"):
            konta = generate_zpk(zpk_branza, zpk_cit, zpk_rzis, zpk_skala,
                                  zpk_aut == "Tak", zpk_tp == "Tak")
            st.session_state["zpk_konta"] = konta
            st.success(f"Wygenerowano {len(konta)} kont!")

        if "zpk_konta" in st.session_state:
            konta = st.session_state["zpk_konta"]
            st.markdown(f"**Wygenerowany plan: {len(konta)} kont**")

            # Podglad
            import pandas as pd
            df = pd.DataFrame(konta)
            st.dataframe(df, use_container_width=True, height=400)

            # Eksport XLSX
            xlsx_buf = zpk_to_xlsx(konta)
            if xlsx_buf:
                st.download_button("Pobierz ZPK jako XLSX", xlsx_buf,
                    f"ZPK_{(G('d_name') or 'spolka').replace(' ','_')}.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True)

    st.markdown("---")
    st.markdown("**System informatyczny**")
    st.session_state.d_sn = st.text_input("Oprogramowanie", value=G("d_sn"), key="wsn", placeholder="np. Symfonia, Enova365")
    c1, c2 = st.columns(2)
    with c1:
        st.session_state.d_sv = st.text_input("Wersja", value=G("d_sv"), key="wsv")
    with c2:
        st.session_state.d_sp = st.text_input("Producent", value=G("d_sp"), key="wsp")

    st.markdown("---")
    st.markdown("### Krajowy System e-Faktur (KSeF)")
    st.caption("Od 2026 r. KSeF jest obowiazkowy. Okresl zasady obiegu faktur ustrukturyzowanych.")

    st.session_state.d_ksef = st.checkbox("Jednostka korzysta z KSeF (obowiazkowe od 2026)",
                                           value=G("d_ksef"), key="wksef")

    if st.session_state.d_ksef:
        st.session_state.d_ksef_moment = st.radio(
            "Moment ujecia faktury w ksiegach",
            ["Data wystawienia w KSeF",
             "Data operacji gospodarczej (dostawa/usluga)",
             "Data otrzymania faktury w KSeF"],
            key="wksef_mom",
            help="Okresl, ktora data jest podstawa ujecia faktury w ksiegach rachunkowych."
        )

        st.session_state.d_ksef_korekty = st.radio(
            "Obsługa faktur korygujacych",
            ["Nota korygujaca w KSeF",
             "Faktura korygujaca z odniesieniem do faktury pierwotnej",
             "Korekta zbiorcza za okres rozliczeniowy"],
            key="wksef_kor"
        )

        st.session_state.d_ksef_archiwum = st.radio(
            "Archiwizacja faktur",
            ["Wylacznie w KSeF (repozytorium MF spelnia wymog archiwizacji)",
             "Rownolegle w KSeF i w systemie FK jednostki",
             "Rownolegle w KSeF, systemie FK i na nosnikach zapasowych"],
            key="wksef_arch"
        )

        st.session_state.d_ksef_system = st.radio(
            "Integracja KSeF z systemem ksiegowym",
            ["Zintegrowany z systemem FK (automatyczny import/eksport)",
             "Polautomatyczny (import plikow XML z KSeF)",
             "Reczny (wprowadzanie na podstawie faktur z KSeF)"],
            key="wksef_sys"
        )


def step_2():
    st.subheader("Krok 3: Metody wyceny")
    st.session_state.d_dep = st.radio("Amortyzacja ST", ["Metoda liniowa", "Metoda degresywna", "Jednorazowa"], key="wdep")
    st.session_state.d_thr = st.slider("Prog ST (PLN)", 3500, 30000, G("d_thr"), 500, key="wthr")
    st.session_state.d_iv = st.radio("Wycena zapasow", ["Cena nabycia", "Koszt wytworzenia", "Cena rynkowa"], key="wiv")
    st.info("**Art. 34 ust. 4 UoR** - wybierz metode rozchodu i stosuj konsekwentnie.")
    st.session_state.d_id = st.radio("Rozchod zapasow", ["FIFO", "LIFO", "Srednia wazona", "Szczegolowa identyfikacja"], key="wid")


def step_3():
    st.subheader("Krok 4: Koszty i RZiS")
    st.session_state.d_cm = st.radio("Model kosztow", ["Tylko Zespol 4 (uklad rodzajowy)", "Tylko Zespol 5 (uklad kalkulacyjny)", "Zespol 4 + 5 (oba uklady)"], key="wcm")
    cm = st.session_state.d_cm
    if "Zespol 4" in cm and "5" not in cm:
        st.session_state.d_pl = "Wariant porownawczy"; st.info("RZiS: **porownawczy** (auto)")
    elif "Zespol 5" in cm and "4" not in cm:
        st.session_state.d_pl = "Wariant kalkulacyjny"; st.info("RZiS: **kalkulacyjny** (auto)")
    else:
        st.session_state.d_pl = st.radio("Wariant RZiS", ["Wariant porownawczy", "Wariant kalkulacyjny"], key="wpl")
    if "Zespol 5" in cm or "4 + 5" in cm:
        st.session_state.d_pc = st.radio("Kalkulacja kosztu", ["Pelny koszt wytworzenia", "Zmienny koszt wytworzenia"], key="wpc")
        st.session_state.d_oh = st.radio("Klucz kosztow posrednich", ["Klucz przychodowy", "Klucz kosztowy", "Bezposrednie przypisanie"], key="woh")


def step_4():
    st.subheader("Krok 5: Waluty obce")
    st.session_state.d_fxs = st.radio("Kurs walutowy", ["Kurs sredni NBP", "Kurs kupna banku", "Kurs sprzedazy banku"], key="wfxs")
    st.session_state.d_hfx = st.checkbox("Rachunki walutowe", value=G("d_hfx"), key="whfx")
    if st.session_state.d_hfx:
        st.session_state.d_fxd = st.radio("Rozchod waluty", ["FIFO", "LIFO", "Srednia wazona"], key="wfxd")
        st.session_state.d_cur = st.multiselect("Waluty", ALL_CUR, default=G("d_cur"), key="wcur")


def step_5():
    st.subheader("Krok 6: Ochrona danych")
    st.session_state.d_dp = st.radio("Metoda ochrony", ["Elektroniczna i fizyczna", "Wylacznie elektroniczna", "Wylacznie fizyczna"], key="wdp")
    st.session_state.d_ay = st.slider("Archiwizacja (lata)", 5, 15, G("d_ay"), key="way")
    st.session_state.d_bk = st.radio("Kopie zapasowe", ["Codziennie", "Co tydzien", "Co miesiac"], key="wbk")
    st.session_state.d_ac = st.checkbox("Kontrola dostepu z haslami", value=G("d_ac"), key="wac")
    st.session_state.d_rp = st.text_input("Osoba odpowiedzialna", value=G("d_rp"), key="wrp")


def step_6():
    st.subheader("Krok 7: Polityki dodatkowe")
    st.session_state.d_rev = st.radio("Przychody", ["Zasada memorialowa", "Zasada kasowa"], key="wrev")
    st.session_state.d_ls = st.radio("Leasing", ["Wg przepisow bilansowych", "Wg przepisow podatkowych"], key="wls")
    st.session_state.d_prov = st.checkbox("Rezerwy (art. 35d)", value=G("d_prov"), key="wprov")
    st.session_state.d_dt = st.checkbox("Podatek odroczony", value=G("d_dt"), key="wdt")
    if not (G("d_small") or G("d_micro")):
        st.session_state.d_cf = st.radio("Przeplywy pieniezne", ["Metoda posrednia", "Metoda bezposrednia"], key="wcf")
    st.markdown("**Zatwierdzenie**")
    c1, c2 = st.columns(2)
    with c1: st.session_state.d_adate = st.date_input("Data zatwierdzenia", value=G("d_adate"), key="wad")
    with c2: st.session_state.d_edate = st.date_input("Data wejscia w zycie", value=G("d_edate"), key="wed")
    st.session_state.d_ab = st.text_input("Zatwierdzil(a)", value=G("d_ab"), key="wab", placeholder="Imie i nazwisko")


def step_7():
    st.subheader("Krok 8: Eksport DOCX")
    buf = gen_docx()
    fn = f"Polityka_Rachunkowosci_{(G('d_name') or 'jednostka').replace(' ', '_')}.docx"
    st.download_button("Pobierz Polityke Rachunkowosci (DOCX)", buf, fn,
                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                       use_container_width=True, type="primary")

    if "zpk_konta" in st.session_state:
        xlsx_buf = zpk_to_xlsx(st.session_state["zpk_konta"])
        if xlsx_buf:
            st.download_button("Pobierz Zakladowy Plan Kont (XLSX)", xlsx_buf,
                f"ZPK_{(G('d_name') or 'spolka').replace(' ', '_')}.xlsx",
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True)

    st.success("Dokumenty gotowe do pobrania!")
    st.divider()
    efi = G("d_form")
    efl = ENTITY_FORM_LABELS[efi] if isinstance(efi, int) and efi < len(ENTITY_FORM_LABELS) else ""
    with st.expander("Podglad danych", expanded=True):
        st.write(f"**{G('d_name') or '-'}** ({efl})")
        st.write(f"NIP: {G('d_nip') or '-'} | KRS: {G('d_krs') or '-'} | REGON: {G('d_regon') or '-'}")
        st.write(f"Koszty: {G('d_cm')} | RZiS: {G('d_pl')}")
        st.write(f"Zatwierdzil(a): {G('d_ab') or '-'}")


# ══════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════

STEPS = [step_0, step_1, step_2, step_3, step_4, step_5, step_6, step_7]
st.title("Generator Polityki Rachunkowosci")
st.caption("Zgodna z Ustawa o Rachunkowosci (art. 10 UoR) | Stan prawny 2026")

prog = st.session_state.step / max(len(STEPS) - 1, 1)
st.progress(prog, text=f"**{STEP_NAMES[st.session_state.step]}** ({st.session_state.step+1}/{len(STEPS)})")

STEPS[st.session_state.step]()

st.divider()
c1, c2, c3 = st.columns([1, 2, 1])
with c1:
    if st.session_state.step > 0 and st.button("Wstecz", use_container_width=True, key="bk"):
        st.session_state.step -= 1; st.rerun()
with c2:
    st.markdown(f"<p style='text-align:center;color:#999;margin-top:8px'>{st.session_state.step+1} / {len(STEPS)}</p>", unsafe_allow_html=True)
with c3:
    if st.session_state.step < len(STEPS) - 1 and st.button("Dalej", use_container_width=True, type="primary", key="fw"):
        st.session_state.step += 1; st.rerun()
